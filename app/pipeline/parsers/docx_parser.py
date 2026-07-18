from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def read_docx(path: Path) -> str:
    if path.suffix.lower() == ".doc":
        raise ValueError(
            f"'{path.name}' is a legacy .doc file (old binary Word format), which "
            "python-docx cannot read. Please re-save it as .docx from Word, "
            "LibreOffice, or Pages first, then try again."
        )

    document = docx.Document(str(path))

    # document.paragraphs / document.tables are two separate flat lists with
    # no interleaving information — collecting one after the other (the
    # previous approach) silently moves every table to the end of the
    # transcript, regardless of where it actually sits in the source
    # document. Walking the body's direct XML children instead preserves
    # true document order, matching how odf_parser.py already handles this
    # for .odt. This matters because a table sitting mid-document (e.g. a
    # stammdaten/enrollment table) would otherwise get relocated to the very
    # end of the transcript — and tabular "|"-joined text is already known
    # to be harder for downstream PII detection than surrounding prose (see
    # the ingest step's tabular-NER caveat in CLAUDE.md), so relocating it
    # concentrates detection misses at the tail of every such document.
    parts = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            parts.append(Paragraph(child, document).text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts)
